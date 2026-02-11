# utils/file_importer.py - Universal file loader: .docx, .pdf, .xlsx, .xls, .csv, .txt, .md
from typing import Tuple, Optional

# Optional deps: import inside methods to avoid hard fail if not installed


def _clean_text(text: str) -> str:
    """Làm sạch whitespace thừa (giữ xuống dòng hợp lý)."""
    if not text or not isinstance(text, str):
        return ""
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


class UniversalLoader:
    """Đọc file dữ liệu (PDF, TXT, DOCX, CSV, XLS, XLSX, MD) và chuyển thành text content."""

    SUPPORTED_EXTENSIONS = (".docx", ".pdf", ".xlsx", ".xls", ".csv", ".txt", ".md")

    @classmethod
    def load(cls, file) -> Tuple[str, Optional[str]]:
        """
        Input: file object (từ st.file_uploader).
        Output: (text đã làm sạch, None) hoặc ("", "thông báo lỗi thân thiện").
        """
        if file is None:
            return "", "Không có file được chọn."

        name = getattr(file, "name", "") or ""
        ext = "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""

        if ext not in cls.SUPPORTED_EXTENSIONS:
            return "", f"Định dạng chưa hỗ trợ. Dùng: {', '.join(cls.SUPPORTED_EXTENSIONS)}"

        try:
            file.seek(0)
            raw = file.read()
        except Exception as e:
            return "", f"Không đọc được file: {str(e)}"

        if ext in (".txt", ".md"):
            return cls._load_text(raw), None
        if ext == ".docx":
            return cls._load_docx(raw, name)
        if ext == ".pdf":
            return cls._load_pdf(raw, name)
        if ext in (".xlsx", ".xls"):
            return cls._load_xlsx(raw, name)
        if ext == ".csv":
            return cls._load_csv(raw, name)

        return "", "Định dạng file chưa được xử lý."

    @classmethod
    def _load_text(cls, raw: bytes) -> str:
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = raw.decode("utf-8", errors="replace")
            except Exception:
                return ""
        return _clean_text(text)

    @classmethod
    def _load_docx(cls, raw: bytes, name: str) -> Tuple[str, Optional[str]]:
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(raw))
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
            return _clean_text(text), None
        except ImportError:
            return "", "Cần cài đặt python-docx: pip install python-docx"
        except Exception as e:
            return "", f"File Word lỗi hoặc không đọc được: {str(e)}"

    @classmethod
    def _load_pdf(cls, raw: bytes, name: str) -> Tuple[str, Optional[str]]:
        try:
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(raw))
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            text = "\n".join(parts)
            return _clean_text(text), None
        except ImportError:
            return "", "Cần cài đặt pypdf: pip install pypdf"
        except Exception as e:
            return "", f"File PDF lỗi hoặc không đọc được: {str(e)}"

    @classmethod
    def _load_xlsx(cls, raw: bytes, name: str) -> Tuple[str, Optional[str]]:
        try:
            import pandas as pd
            from io import BytesIO
            df = pd.read_excel(BytesIO(raw), sheet_name=None, header=None)
            parts = []
            for sheet_name, frame in df.items():
                parts.append(f"[Sheet: {sheet_name}]")
                parts.append(frame.to_string(index=False, header=False))
            text = "\n\n".join(parts)
            return _clean_text(text), None
        except ImportError:
            return "", "Cần cài đặt pandas và openpyxl: pip install pandas openpyxl"
        except Exception as e:
            return "", f"File Excel lỗi hoặc không đọc được: {str(e)}"

    @classmethod
    def _load_csv(cls, raw: bytes, name: str) -> Tuple[str, Optional[str]]:
        try:
            import pandas as pd
            from io import BytesIO
            try:
                df = pd.read_csv(BytesIO(raw), encoding="utf-8", on_bad_lines="skip")
            except TypeError:
                df = pd.read_csv(BytesIO(raw), encoding="utf-8", error_bad_lines=False)
            text = df.to_string(index=False)
            return _clean_text(text), None
        except UnicodeDecodeError:
            try:
                df = pd.read_csv(BytesIO(raw), encoding="utf-8", errors="replace")
                text = df.to_string(index=False)
                return _clean_text(text), None
            except Exception as e:
                return "", f"File CSV lỗi hoặc encoding không đúng: {str(e)}"
        except ImportError:
            return "", "Cần cài đặt pandas: pip install pandas"
        except Exception as e:
            return "", f"File CSV lỗi hoặc không đọc được: {str(e)}"
